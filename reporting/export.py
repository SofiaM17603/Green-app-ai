"""
Export carbon reports to PDF and DOCX formats
"""

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Image, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.linecharts import HorizontalLineChart

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import io
from datetime import datetime
from typing import Dict
from .templates.text_templates import get_template, format_number


class PDFExporter:
    """Export carbon report to PDF"""

    def __init__(self, report_data: Dict, output_path: str):
        self.report_data = report_data
        self.output_path = output_path
        self.lang = report_data['metadata']['language']
        self.t = get_template(self.lang)

        # Colors
        self.color_primary = colors.HexColor('#047857')  # Forest green
        self.color_secondary = colors.HexColor('#10b981')  # Emerald
        self.color_accent = colors.HexColor('#fbbf24')  # Gold
        self.color_text = colors.HexColor('#0a1f1a')  # Dark green

    def create_header_footer(self, canvas_obj, doc):
        """Add header and footer to each page"""
        canvas_obj.saveState()

        # Header
        canvas_obj.setFillColor(self.color_primary)
        canvas_obj.rect(0, A4[1] - 1.5*cm, A4[0], 1.5*cm, fill=True, stroke=False)

        canvas_obj.setFillColor(colors.white)
        canvas_obj.setFont('Helvetica-Bold', 16)
        canvas_obj.drawString(2*cm, A4[1] - 1*cm, '🌱 Green App')

        canvas_obj.setFont('Helvetica', 10)
        canvas_obj.drawRightString(A4[0] - 2*cm, A4[1] - 1*cm, self.t['report_title'])

        # Footer
        canvas_obj.setFillColor(colors.grey)
        canvas_obj.setFont('Helvetica', 8)
        canvas_obj.drawString(
            2*cm,
            1*cm,
            f"{self.t['generated_on']} {self.report_data['metadata']['generated_date']}"
        )
        canvas_obj.drawRightString(
            A4[0] - 2*cm,
            1*cm,
            f"{self.t['page']} {doc.page}"
        )

        canvas_obj.restoreState()

    def generate(self):
        """Generate PDF report"""
        doc = SimpleDocTemplate(
            self.output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2.5*cm,
            bottomMargin=2*cm
        )

        # Build story
        story = []
        styles = self._create_styles()

        # Title page
        story.extend(self._create_title_page(styles))

        # Executive summary
        story.extend(self._create_executive_summary(styles))
        story.append(PageBreak())

        # Emissions breakdown
        story.extend(self._create_emissions_breakdown(styles))
        story.append(PageBreak())

        # Evolution
        story.extend(self._create_evolution_section(styles))
        story.append(PageBreak())

        # Recommendations
        story.extend(self._create_recommendations(styles))
        story.append(PageBreak())

        # Methodology
        story.extend(self._create_methodology(styles))

        # Build PDF
        doc.build(story, onFirstPage=self.create_header_footer, onLaterPages=self.create_header_footer)

    def _create_styles(self):
        """Create custom paragraph styles"""
        styles = getSampleStyleSheet()

        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.color_primary,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))

        styles.add(ParagraphStyle(
            name='SectionHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=self.color_primary,
            spaceBefore=20,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        ))

        styles.add(ParagraphStyle(
            name='SubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=self.color_text,
            spaceBefore=10,
            spaceAfter=8,
            fontName='Helvetica-Bold'
        ))

        styles.add(ParagraphStyle(
            name='CustomBodyText',
            parent=styles['Normal'],
            fontSize=10,
            textColor=self.color_text,
            alignment=TA_JUSTIFY,
            spaceBefore=6,
            spaceAfter=6
        ))

        styles.add(ParagraphStyle(
            name='CustomHighlight',
            parent=styles['Normal'],
            fontSize=14,
            textColor=self.color_primary,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceBefore=10,
            spaceAfter=10
        ))

        return styles

    def _create_title_page(self, styles):
        """Create title page"""
        elements = []

        elements.append(Spacer(1, 2*cm))

        # Main title
        title = Paragraph(self.t['report_title'], styles['CustomTitle'])
        elements.append(title)

        elements.append(Spacer(1, 1*cm))

        # Subtitle
        period = self.report_data['summary']['period']
        if period['start_date'] and period['end_date']:
            subtitle = f"{period['start_date']} - {period['end_date']}"
        else:
            subtitle = self.report_data['metadata']['generated_date']

        elements.append(Paragraph(subtitle, styles['Heading2']))

        elements.append(Spacer(1, 3*cm))

        # Key figures box
        total = format_number(
            self.report_data['summary']['total_emissions_tons'],
            self.lang
        )

        key_fig = f"""
        <para align=center>
            <font size=40 color={self.color_primary}><b>{total}</b></font><br/>
            <font size=14>{self.t['unit_tons']}</font>
        </para>
        """
        elements.append(Paragraph(key_fig, styles['Normal']))

        elements.append(Spacer(1, 1*cm))

        # Standards
        standards_text = f"<para align=center><i>{self.t['standards_compliance']}:<br/>ADEME Bilan Carbone® | GHG Protocol</i></para>"
        elements.append(Paragraph(standards_text, styles['CustomBodyText']))

        elements.append(PageBreak())

        return elements

    def _create_executive_summary(self, styles):
        """Create executive summary section"""
        elements = []

        elements.append(Paragraph(self.t['executive_summary'], styles['SectionHeading']))
        elements.append(Spacer(1, 0.5*cm))

        summary = self.report_data['summary']

        # Create summary table
        data = [
            [self.t['total_emissions'], f"{format_number(summary['total_emissions_tons'], self.lang)} {self.t['unit_tons']}"],
            [self.t['invoice_analyzed'], f"{summary['invoice_count']}"],
            [self.t['average_per_invoice'], f"{format_number(summary['average_per_invoice'], self.lang)} {self.t['unit_co2e']}"],
        ]

        if summary['period']['start_date']:
            data.append([
                self.t['reporting_period'],
                f"{summary['period']['start_date']} - {summary['period']['end_date']}"
            ])

        table = Table(data, colWidths=[8*cm, 8*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), self.color_primary),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))

        elements.append(table)
        elements.append(Spacer(1, 1*cm))

        # Pie chart - emissions by scope
        if self.report_data['breakdown']['by_scope']:
            chart_img = self._create_pie_chart(
                self.report_data['breakdown']['by_scope'],
                self.t['by_scope']
            )
            elements.append(chart_img)

        return elements

    def _create_emissions_breakdown(self, styles):
        """Create emissions breakdown section"""
        elements = []

        elements.append(Paragraph(self.t['emissions_breakdown'], styles['SectionHeading']))
        elements.append(Spacer(1, 0.5*cm))

        # By category
        elements.append(Paragraph(self.t['by_category'], styles['SubHeading']))

        by_category = self.report_data['breakdown']['by_category']
        if by_category:
            data = [['Catégorie' if self.lang == 'fr' else 'Category', 'CO₂e (kg)', 'Count', '%']]

            total = sum(cat['total'] for cat in by_category.values())

            for category, values in sorted(by_category.items(), key=lambda x: x[1]['total'], reverse=True):
                pct = (values['total'] / total * 100) if total > 0 else 0
                data.append([
                    category,
                    format_number(values['total'], self.lang),
                    str(values['count']),
                    f"{pct:.1f}%"
                ])

            table = Table(data, colWidths=[6*cm, 4*cm, 3*cm, 3*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.color_primary),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))

            elements.append(table)

        return elements

    def _create_evolution_section(self, styles):
        """Create evolution section with chart"""
        elements = []

        elements.append(Paragraph(self.t['evolution'], styles['SectionHeading']))
        elements.append(Spacer(1, 0.5*cm))

        monthly = self.report_data['evolution']['monthly']
        if monthly:
            chart_img = self._create_line_chart(monthly, self.t['monthly_trend'])
            elements.append(chart_img)

        return elements

    def _create_recommendations(self, styles):
        """Create recommendations section"""
        elements = []

        elements.append(Paragraph(self.t['recommendations'], styles['SectionHeading']))
        elements.append(Spacer(1, 0.5*cm))

        for rec in self.report_data['recommendations']:
            # Priority badge
            priority_color = {
                'high': self.color_accent,
                'medium': self.color_secondary,
                'low': colors.grey
            }.get(rec['priority'], colors.grey)

            elements.append(Paragraph(
                f"<b><font color={priority_color}>⬤</font> {rec['title']}</b>",
                styles['SubHeading']
            ))

            elements.append(Paragraph(rec['description'], styles['CustomBodyText']))

            if 'potential_reduction' in rec:
                elements.append(Paragraph(
                    f"<i>{self.t['potential_reduction']}: {format_number(rec['potential_reduction'], self.lang)} kg CO₂e</i>",
                    styles['CustomBodyText']
                ))

            elements.append(Spacer(1, 0.3*cm))

        # Climate commitments
        elements.append(Spacer(1, 1*cm))
        elements.append(Paragraph(self.t['commitments'], styles['SectionHeading']))
        elements.append(Paragraph(
            self.report_data['climate_commitments'],
            styles['CustomBodyText']
        ))

        return elements

    def _create_methodology(self, styles):
        """Create methodology section"""
        elements = []

        methodology = self.report_data['methodology']

        elements.append(Paragraph(methodology['title'], styles['SectionHeading']))
        elements.append(Paragraph(methodology['text'], styles['CustomBodyText']))

        elements.append(Spacer(1, 0.5*cm))
        elements.append(Paragraph("<b>Sources:</b>", styles['SubHeading']))

        for source in methodology['sources']:
            elements.append(Paragraph(f"• {source}", styles['CustomBodyText']))

        return elements

    def _create_pie_chart(self, data: dict, title: str):
        """Create pie chart using matplotlib"""
        fig, ax = plt.subplots(figsize=(6, 4))

        labels = list(data.keys())
        values = list(data.values())

        colors_list = ['#047857', '#10b981', '#34d399', '#6ee7b7', '#a7f3d0']

        ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors_list)
        ax.set_title(title, fontsize=12, fontweight='bold')

        # Save to BytesIO
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()

        # Create Image for ReportLab
        img = Image(img_buffer, width=12*cm, height=8*cm)
        return img

    def _create_line_chart(self, data: dict, title: str):
        """Create line chart using matplotlib"""
        fig, ax = plt.subplots(figsize=(10, 5))

        months = list(data.keys())
        values = list(data.values())

        ax.plot(months, values, marker='o', linewidth=2, color='#047857')
        ax.fill_between(range(len(values)), values, alpha=0.3, color='#10b981')

        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel('Période' if self.lang == 'fr' else 'Period')
        ax.set_ylabel('CO₂e (kg)')
        ax.grid(True, alpha=0.3)

        plt.xticks(rotation=45)
        plt.tight_layout()

        # Save to BytesIO
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()

        img = Image(img_buffer, width=16*cm, height=8*cm)
        return img


class DOCXExporter:
    """Export carbon report to DOCX"""

    def __init__(self, report_data: Dict, output_path: str):
        self.report_data = report_data
        self.output_path = output_path
        self.lang = report_data['metadata']['language']
        self.t = get_template(self.lang)
        self.doc = Document()

        # Set styles
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles"""
        styles = self.doc.styles

        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(4, 120, 87)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(30)

    def generate(self):
        """Generate DOCX report"""
        # Title page
        self._create_title_page()
        self.doc.add_page_break()

        # Executive summary
        self._create_executive_summary()
        self.doc.add_page_break()

        # Emissions breakdown
        self._create_emissions_breakdown()
        self.doc.add_page_break()

        # Recommendations
        self._create_recommendations()
        self.doc.add_page_break()

        # Methodology
        self._create_methodology()

        # Save
        self.doc.save(self.output_path)

    def _create_title_page(self):
        """Create title page"""
        # Title
        title = self.doc.add_heading(self.t['report_title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Period
        period = self.report_data['summary']['period']
        if period['start_date']:
            subtitle = f"{period['start_date']} - {period['end_date']}"
        else:
            subtitle = self.report_data['metadata']['generated_date']

        p = self.doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Key figure
        self.doc.add_paragraph()
        total = format_number(self.report_data['summary']['total_emissions_tons'], self.lang)

        p = self.doc.add_paragraph()
        run = p.add_run(f"{total} {self.t['unit_tons']}")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(4, 120, 87)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _create_executive_summary(self):
        """Create executive summary"""
        self.doc.add_heading(self.t['executive_summary'], level=1)

        summary = self.report_data['summary']

        # Table
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        cells = table.rows[0].cells
        cells[0].text = self.t['total_emissions']
        cells[1].text = f"{format_number(summary['total_emissions_tons'], self.lang)} {self.t['unit_tons']}"

        cells = table.rows[1].cells
        cells[0].text = self.t['invoice_analyzed']
        cells[1].text = str(summary['invoice_count'])

        cells = table.rows[2].cells
        cells[0].text = self.t['average_per_invoice']
        cells[1].text = f"{format_number(summary['average_per_invoice'], self.lang)} {self.t['unit_co2e']}"

        if summary['period']['start_date']:
            cells = table.rows[3].cells
            cells[0].text = self.t['reporting_period']
            cells[1].text = f"{summary['period']['start_date']} - {summary['period']['end_date']}"

    def _create_emissions_breakdown(self):
        """Create emissions breakdown"""
        self.doc.add_heading(self.t['emissions_breakdown'], level=1)

        by_category = self.report_data['breakdown']['by_category']
        if by_category:
            table = self.doc.add_table(rows=len(by_category)+1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Catégorie' if self.lang == 'fr' else 'Category'
            header_cells[1].text = 'CO₂e (kg)'
            header_cells[2].text = 'Count'
            header_cells[3].text = '%'

            total = sum(cat['total'] for cat in by_category.values())

            # Data rows
            for idx, (category, values) in enumerate(sorted(by_category.items(), key=lambda x: x[1]['total'], reverse=True), 1):
                pct = (values['total'] / total * 100) if total > 0 else 0
                cells = table.rows[idx].cells
                cells[0].text = category
                cells[1].text = format_number(values['total'], self.lang)
                cells[2].text = str(values['count'])
                cells[3].text = f"{pct:.1f}%"

    def _create_recommendations(self):
        """Create recommendations section"""
        self.doc.add_heading(self.t['recommendations'], level=1)

        for rec in self.report_data['recommendations']:
            self.doc.add_heading(rec['title'], level=2)
            self.doc.add_paragraph(rec['description'])

            if 'potential_reduction' in rec:
                p = self.doc.add_paragraph()
                p.add_run(f"{self.t['potential_reduction']}: ").italic = True
                p.add_run(f"{format_number(rec['potential_reduction'], self.lang)} kg CO₂e").bold = True

        # Climate commitments
        self.doc.add_heading(self.t['commitments'], level=1)
        self.doc.add_paragraph(self.report_data['climate_commitments'])

    def _create_methodology(self):
        """Create methodology section"""
        methodology = self.report_data['methodology']

        self.doc.add_heading(methodology['title'], level=1)
        self.doc.add_paragraph(methodology['text'])

        self.doc.add_heading('Sources:', level=2)
        for source in methodology['sources']:
            self.doc.add_paragraph(source, style='List Bullet')


def export_pdf(report_data: Dict, output_path: str):
    """Export report to PDF"""
    exporter = PDFExporter(report_data, output_path)
    exporter.generate()
    return output_path


def export_docx(report_data: Dict, output_path: str):
    """Export report to DOCX"""
    exporter = DOCXExporter(report_data, output_path)
    exporter.generate()
    return output_path
